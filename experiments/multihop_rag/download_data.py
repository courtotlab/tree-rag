#!/usr/bin/env python3
"""Download, sample, and materialize the public ODC-BY MultiHop-RAG corpus."""

from __future__ import annotations

import json
import random
import re
import urllib.request
from collections import defaultdict
from pathlib import Path

from docx import Document


HERE = Path(__file__).resolve().parent
INPUT = HERE / "input"
FOLDERS = HERE / "folders"
BASE = "https://huggingface.co/datasets/yixuantt/MultiHopRAG/resolve/main"
SEED = 20260806


def download(name: str) -> Path:
    INPUT.mkdir(exist_ok=True)
    path = INPUT / name
    if not path.exists():
        print(f"downloading {name}")
        urllib.request.urlretrieve(f"{BASE}/{name}", path)
    return path


def safe(value: str, limit: int = 80) -> str:
    value = re.sub(r"[^\w\s\-.,()]", "", value or "untitled").strip()
    return re.sub(r"\s+", " ", value)[:limit].strip() or "untitled"


def materialize(corpus: list[dict]) -> None:
    seen = set()
    for row in corpus:
        category, source = safe(row.get("category", "misc"), 40), safe(row.get("source", "unknown"), 60)
        title = safe(row.get("title", "untitled"))
        base_title, suffix = title, 2
        while (category, source, title) in seen:
            title = f"{base_title} ({suffix})"
            suffix += 1
        seen.add((category, source, title))
        target = FOLDERS / category / source
        target.mkdir(parents=True, exist_ok=True)
        path = target / f"{title}.docx"
        if path.exists():
            continue
        doc = Document()
        doc.add_heading(row.get("title") or "Untitled", level=1)
        doc.add_paragraph(
            f"Source: {row.get('source','')} | Category: {row.get('category','')} | "
            f"Published: {row.get('published_at','')} | Author: {row.get('author') or 'unknown'}"
        )
        for paragraph in (row.get("body") or "").split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        doc.save(path)


def sample_queries(queries: list[dict]) -> list[dict]:
    groups = defaultdict(list)
    for index, row in enumerate(queries):
        groups[row["question_type"]].append(index)
    rng = random.Random(SEED)
    indices = sorted(index for values in groups.values() for index in rng.sample(values, 50))
    return [queries[index] | {"qid": f"MHR-{index:04d}"} for index in indices]


def main() -> None:
    corpus = json.loads(download("corpus.json").read_text(encoding="utf-8"))
    queries = json.loads(download("MultiHopRAG.json").read_text(encoding="utf-8"))
    materialize(corpus)
    sample = sample_queries(queries)
    (INPUT / "sample_200.json").write_text(json.dumps(sample, indent=2), encoding="utf-8")
    print(f"materialized {len(corpus)} documents; sampled {len(sample)} questions")


if __name__ == "__main__":
    main()

